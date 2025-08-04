import os
import io
import traceback
import re
import json
import requests
from flask import Flask, request, send_file, render_template
from flask_cors import CORS

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from weasyprint import HTML, CSS
import arabic_reshaper
from bidi.algorithm import get_display

app = Flask(__name__)
CORS(app)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FONT_FILE_NAME = "Vazirmatn-Regular.ttf"
FOOTER_TEXT = "هوش مصنوعی آلفا دانلود از گوگل پلی"

# --- توابع کمکی ---
def reshape_text(text):
    if not text:
        return ""
    return get_display(arabic_reshaper.reshape(text))

# --- توابع جدید برای ساخت فایل از داده‌های ساختاریافته (JSON) ---

def create_html_from_data(data):
    """HTML را از روی داده‌های ساختاریافته چت (شامل تصاویر) تولید می‌کند."""
    title = data.get('title', 'گفتگو')
    messages = data.get('messages', [])
    
    html_parts = []
    for msg in messages:
        role = "کاربر" if msg.get('role') == 'user' else "مدل"
        text = msg.get('text', '')
        images = msg.get('images', [])
        
        # استایل بر اساس نقش (کاربر یا مدل)
        style = "background-color: #e0f0ff; border-right: 3px solid #007bff;" if role == "کاربر" else "background-color: #f8f9fa; border-right: 3px solid #6c757d;"
        
        html_parts.append(f'<div class="message" style="{style}">')
        html_parts.append(f'<strong>{reshape_text(role)}:</strong>')
        
        if text:
            # برای نمایش صحیح پاراگراف‌ها در HTML
            reshaped_text_content = reshape_text(text).replace('\n', '<br>')
            html_parts.append(f'<p>{reshaped_text_content}</p>')
            
        if images:
            html_parts.append('<div class="image-container">')
            for img_url in images:
                html_parts.append(f'<img src="{img_url}" alt="Image from chat">')
            html_parts.append('</div>')
            
        html_parts.append('</div>')
        
    chat_html = "\n".join(html_parts)
    reshaped_footer = reshape_text(FOOTER_TEXT)
    
    full_html = f"""
    <!DOCTYPE html>
    <html lang="fa" dir="rtl">
    <head>
        <meta charset="UTF-8">
        <title>{reshape_text(title)}</title>
        <style>
            @font-face {{ font-family: 'Vazir'; src: url('{FONT_FILE_NAME}'); }}
            body {{ font-family: 'Vazir', sans-serif; font-size: 12pt; line-height: 1.8; max-width: 800px; margin: 2rem auto; padding: 2rem; border: 1px solid #ddd; }}
            h1 {{ text-align: center; border-bottom: 2px solid #eee; padding-bottom: 10px; }}
            .message {{ padding: 15px; margin-bottom: 15px; border-radius: 8px; }}
            .message p {{ margin-top: 5px; white-space: pre-wrap; }}
            .image-container {{ margin-top: 10px; }}
            .image-container img {{ max-width: 100%; height: auto; border-radius: 5px; margin-bottom: 10px; border: 1px solid #ccc; }}
            .footer {{ margin-top: 2rem; padding-top: 1rem; border-top: 1px solid #eee; text-align: center; color: #007bff; font-size: 10pt; }}
        </style>
    </head>
    <body>
        <h1>{reshape_text(title)}</h1>
        {chat_html}
        <div class="footer">{reshaped_footer}</div>
    </body>
    </html>
    """
    return full_html

def create_pdf_from_data(data):
    """PDF را از داده‌های ساختاریافته تولید می‌کند."""
    full_html = create_html_from_data(data)
    try:
        html = HTML(string=full_html, base_url=BASE_DIR)
        return io.BytesIO(html.write_pdf())
    except Exception:
        print(f"🔥🔥🔥 WEASYPRINT FAILED! 🔥🔥🔥\n{traceback.format_exc()}")
        return io.BytesIO(b"Error generating PDF.")

def create_docx_from_data(data):
    """DOCX را از داده‌های ساختاریافته (شامل دانلود و درج تصاویر) تولید می‌کند."""
    document = Document()
    title = data.get('title', 'گفتگو')
    messages = data.get('messages', [])
    
    # افزودن عنوان
    document.add_heading(reshape_text(title), level=1)

    for msg in messages:
        role = "کاربر" if msg.get('role') == 'user' else "مدل"
        text = msg.get('text', '')
        images = msg.get('images', [])

        # افزودن نقش
        p_role = document.add_paragraph()
        p_role.add_run(reshape_text(role) + ':').bold = True
        
        # افزودن متن
        if text:
            p_text = document.add_paragraph(reshape_text(text))
            p_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
        # افزودن تصاویر
        if images:
            for img_url in images:
                try:
                    # دانلود تصویر از URL
                    response = requests.get(img_url, stream=True)
                    response.raise_for_status()
                    image_stream = io.BytesIO(response.content)
                    document.add_picture(image_stream, width=Inches(4.0))
                except Exception as e:
                    print(f"Error fetching or adding image {img_url}: {e}")
                    document.add_paragraph(f"[خطا در بارگذاری تصویر: {img_url}]")

    # افزودن پاورقی
    footer = document.sections[0].footer
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.text = reshape_text(FOOTER_TEXT)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def create_txt_from_data(data):
    """فایل متنی ساده را از داده‌های ساختاریافته تولید می‌کند."""
    title = data.get('title', 'گفتگو')
    messages = data.get('messages', [])
    
    txt_parts = [f"عنوان: {title}\n{'='*20}\n"]
    for msg in messages:
        role = "کاربر" if msg.get('role') == 'user' else "مدل"
        text = msg.get('text', '')
        images = msg.get('images', [])
        
        txt_parts.append(f"{role}:")
        if text:
            txt_parts.append(text)
        if images:
            for img_url in images:
                txt_parts.append(f"[تصویر: {img_url}]")
        txt_parts.append("\n---\n")
        
    full_content = "\n".join(txt_parts) + f"\n\n{FOOTER_TEXT}"
    return io.BytesIO(full_content.encode('utf-8'))


# --- توابع قدیمی (برای سازگاری با تبدیل پیام تکی) ---
def get_base_html_for_single_message(text_content):
    return f"<p>{reshape_text(text_content).replace('\n', '<br>')}</p>"

def create_pdf_from_string(text_content):
    html_body = get_base_html_for_single_message(text_content)
    full_html = f"<!DOCTYPE html><html lang='fa'><head><meta charset='UTF-8'><style>@font-face {{ font-family: 'Vazir'; src: url('{FONT_FILE_NAME}'); }} body {{ font-family: 'Vazir', sans-serif; }}</style></head><body>{html_body}</body></html>"
    return io.BytesIO(HTML(string=full_html, base_url=BASE_DIR).write_pdf())

def create_docx_from_string(text_content):
    document = Document()
    document.add_paragraph(reshape_text(text_content))
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def create_txt_from_string(text_content):
    return io.BytesIO(text_content.encode('utf-8'))

def create_html_from_string(text_content):
    html_body = get_base_html_for_single_message(text_content)
    return io.BytesIO(f"<html><body>{html_body}</body></html>".encode('utf-8'))


# --- تابع پردازشگر اصلی درخواست ---
@app.route('/', methods=['POST', 'GET', 'HEAD'])
def process_request():
    if request.method == 'HEAD':
        return '', 200
    if request.method == 'GET':
        return render_template('index.html')

    file_format = request.form.get('format', 'txt').lower()
    json_content = request.form.get('json_content')
    
    mimetypes = {'pdf': 'application/pdf', 'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'html': 'text/html', 'txt': 'text/plain'}
    mimetype = mimetypes.get(file_format, 'text/plain')
    filename = f'export.{file_format}'
    
    buffer = None
    
    if json_content:
        # حالت جدید: پردازش JSON
        try:
            data = json.loads(json_content)
            actions = {'pdf': create_pdf_from_data, 'docx': create_docx_from_data, 'html': lambda d: io.BytesIO(create_html_from_data(d).encode('utf-8')), 'txt': create_txt_from_data}
            buffer_func = actions.get(file_format)
            if buffer_func:
                buffer = buffer_func(data)
        except Exception as e:
            print(f"Error processing JSON content: {e}")
            return f"Error processing JSON: {e}", 400
    else:
        # حالت قدیمی: پردازش متن ساده
        content = request.form.get('content')
        if not content:
            return "لطفا محتوایی برای تبدیل وارد کنید.", 400
        actions = {'pdf': create_pdf_from_string, 'docx': create_docx_from_string, 'html': create_html_from_string, 'txt': create_txt_from_string}
        buffer_func = actions.get(file_format)
        if buffer_func:
            buffer = buffer_func(content)

    if buffer:
        return send_file(buffer, as_attachment=True, download_name=filename, mimetype=mimetype)
    else:
        return "فرمت فایل نامعتبر است.", 400

if __name__ == '__main__':
    app.run(debug=True)
